import os
import uuid
from typing import Optional, List, Dict, Any
from datetime import datetime

from app.models.resume import Resume, ExportRequest

class ExportService:
    def __init__(self):
        self.export_directory = "exports"
        os.makedirs(self.export_directory, exist_ok=True)

    async def generate_resume_file(
        self, 
        resume: Resume, 
        export_request: ExportRequest, 
        user_id: str
    ) -> str:
        """Generate resume file in specified format"""

        # Generate unique filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = export_request.filename or f"resume_{timestamp}"

        if export_request.format == "pdf":
            file_path = await self._generate_pdf(resume, filename)
        elif export_request.format == "docx":
            file_path = await self._generate_docx(resume, filename)
        else:
            raise ValueError(f"Unsupported format: {export_request.format}")

        # Record export in database (simplified)
        await self._record_export(user_id, resume.id, file_path, export_request.format)

        return file_path

    async def _generate_pdf(self, resume: Resume, filename: str) -> str:
        """Generate PDF resume"""

        # Simple PDF generation using reportlab
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        file_path = os.path.join(self.export_directory, f"{filename}.pdf")

        doc = SimpleDocTemplate(file_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        # Title
        if resume.content.personal_info:
            title = Paragraph(resume.content.personal_info.name, styles['Title'])
            story.append(title)
            story.append(Spacer(1, 12))

            # Contact info
            contact = f"{resume.content.personal_info.email}"
            if resume.content.personal_info.phone:
                contact += f" | {resume.content.personal_info.phone}"
            contact_para = Paragraph(contact, styles['Normal'])
            story.append(contact_para)
            story.append(Spacer(1, 12))

        # Summary
        if resume.content.summary:
            summary_title = Paragraph("Professional Summary", styles['Heading2'])
            story.append(summary_title)
            summary_para = Paragraph(resume.content.summary, styles['Normal'])
            story.append(summary_para)
            story.append(Spacer(1, 12))

        # Experience
        if resume.content.experience:
            exp_title = Paragraph("Experience", styles['Heading2'])
            story.append(exp_title)

            for exp in resume.content.experience:
                job_title = Paragraph(f"<b>{exp.title}</b> - {exp.company}", styles['Heading3'])
                story.append(job_title)

                date_range = f"{exp.start_date} - {exp.end_date or 'Present'}"
                date_para = Paragraph(date_range, styles['Normal'])
                story.append(date_para)

                for bullet in exp.bullets:
                    bullet_para = Paragraph(f"• {bullet}", styles['Normal'])
                    story.append(bullet_para)

                story.append(Spacer(1, 12))

        # Skills
        if resume.content.skills:
            skills_title = Paragraph("Skills", styles['Heading2'])
            story.append(skills_title)
            skills_text = ", ".join(resume.content.skills)
            skills_para = Paragraph(skills_text, styles['Normal'])
            story.append(skills_para)
            story.append(Spacer(1, 12))

        # Education
        if resume.content.education:
            edu_title = Paragraph("Education", styles['Heading2'])
            story.append(edu_title)

            for edu in resume.content.education:
                edu_para = Paragraph(f"{edu.degree} - {edu.school}", styles['Normal'])
                story.append(edu_para)
                if edu.graduation_year:
                    year_para = Paragraph(edu.graduation_year, styles['Normal'])
                    story.append(year_para)

        doc.build(story)
        return file_path

    async def _generate_docx(self, resume: Resume, filename: str) -> str:
        """Generate DOCX resume"""

        from docx import Document
        from docx.shared import Inches

        file_path = os.path.join(self.export_directory, f"{filename}.docx")

        doc = Document()

        # Title
        if resume.content.personal_info:
            title = doc.add_heading(resume.content.personal_info.name, 0)
            title.alignment = 1  # Center alignment

            # Contact info
            contact = doc.add_paragraph()
            contact.add_run(resume.content.personal_info.email)
            if resume.content.personal_info.phone:
                contact.add_run(f" | {resume.content.personal_info.phone}")
            contact.alignment = 1  # Center alignment

        # Summary
        if resume.content.summary:
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(resume.content.summary)

        # Experience
        if resume.content.experience:
            doc.add_heading('Experience', level=1)

            for exp in resume.content.experience:
                # Job title and company
                job_heading = doc.add_paragraph()
                job_heading.add_run(f"{exp.title} - {exp.company}").bold = True

                # Date range
                date_para = doc.add_paragraph(f"{exp.start_date} - {exp.end_date or 'Present'}")

                # Bullets
                for bullet in exp.bullets:
                    bullet_para = doc.add_paragraph(bullet, style='List Bullet')

        # Skills
        if resume.content.skills:
            doc.add_heading('Skills', level=1)
            skills_text = ", ".join(resume.content.skills)
            doc.add_paragraph(skills_text)

        # Education
        if resume.content.education:
            doc.add_heading('Education', level=1)

            for edu in resume.content.education:
                edu_para = doc.add_paragraph(f"{edu.degree} - {edu.school}")
                if edu.graduation_year:
                    edu_para.add_run(f" ({edu.graduation_year})")

        doc.save(file_path)
        return file_path

    async def _record_export(self, user_id: str, resume_id: str, file_path: str, format: str):
        """Record export in database"""
        # This would typically save to database
        pass

    async def get_user_export_history(self, user_id: str) -> List[Dict[str, Any]]:
        """Get user's export history"""
        # This would typically fetch from database
        return [
            {
                "id": "1",
                "filename": "resume_20241201.pdf",
                "format": "pdf",
                "created_at": "2024-12-01T10:00:00",
                "file_size": "245KB"
            }
        ]
